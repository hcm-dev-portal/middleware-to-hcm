# backend/app/reports/service.py
import os
import re
import uuid
import json
import math
import random
import logging
import tempfile
from enum import Enum
from dataclasses import dataclass, field
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime, timedelta

from fastapi import HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

# Import the enhanced LLM client
from app.reports.llm_client import LLMClient, Intent, Narrative
from app.services.llm.openai_service import OpenAIService
from app.services.aws.translation_service import AWSTranslationService

logger = logging.getLogger(__name__)

# ============================================
# Pydantic models (shared with API signatures)
# ============================================
class ReportAnalysisRequest(BaseModel):
    query: str
    language: str = "en-US"  # Added language support

class ReportGenerationRequest(BaseModel):
    query: str
    analysis: Dict[str, Any]
    format: str = "docx"
    language: str = "en-US"  # Added language support

# In-memory session store (swap for DB/redis in prod)
report_sessions: Dict[str, Dict] = {}

# =============================
# Initialize LLM Client
# =============================
def get_llm_client() -> LLMClient:
    """Factory function to create LLM client with proper services."""
    try:
        openai_service = OpenAIService(model_name="gpt-4o-mini", temperature=0.2)
        translation_service = AWSTranslationService()
        return LLMClient(llm=openai_service, translator=translation_service)
    except Exception as e:
        logger.warning(f"Failed to initialize full LLM client: {e}. Using fallback.")
        # Return basic LLM client that will use fallback methods
        return LLMClient()

# Global LLM client instance
llm_client = get_llm_client()

# =====================
# State machine / plan
# =====================
class ReportType(str, Enum):
    LEAVE = "leave_analysis"
    OVERTIME = "overtime_analysis"
    ATTENDANCE = "attendance_analysis"
    BALANCE = "balance_report"
    UNKNOWN = "unknown"

class Phase(str, Enum):
    COLLECT = "collect_requirements"
    PLAN = "plan"
    GATHER = "gather_data"
    WRITE = "write"
    COMPLETE = "complete"

@dataclass
class PlanStep:
    name: str
    fn: str
    args: Dict[str, Any] = field(default_factory=dict)
    output_key: str = ""

@dataclass
class ExecutionPlan:
    phase: Phase
    steps: List[PlanStep] = field(default_factory=list)
    assumptions: List[str] = field(default_factory=list)

# -----------------------
# Tool/Function registry
# -----------------------
def tool_get_leave_metrics(start: Optional[datetime], end: Optional[datetime], departments: List[str]) -> Dict[str, Any]:
    """Generate realistic leave metrics with some variation."""
    rng = (start.strftime("%b %d") + " – " + end.strftime("%b %d, %Y")) if (start and end) else "Selected period"
    depts = (departments if departments and departments != ["all"] else ["HR", "Engineering", "Sales", "Marketing", "Finance", "Operations"])
    
    # Generate more realistic data with patterns
    base_leave = 40
    dept_data = {}
    for d in depts:
        # Different departments have different patterns
        multiplier = {"HR": 1.2, "Engineering": 0.9, "Sales": 1.3, "Marketing": 1.0, "Finance": 0.8, "Operations": 1.1}.get(d.title(), 1.0)
        dept_data[d.title()] = int(base_leave * multiplier * random.uniform(0.8, 1.5))
    
    total_leave = sum(dept_data.values())
    employee_count = len(depts) * 25  # Assume 25 employees per department
    
    return {
        "range": rng,
        "by_type": {
            "Vacation": int(total_leave * 0.55),
            "Sick": int(total_leave * 0.30),
            "Personal": int(total_leave * 0.10),
            "Bereavement": int(total_leave * 0.05)
        },
        "by_department": dept_data,
        "totals": {
            "leave_days": total_leave,
            "avg_days_per_employee": round(total_leave / employee_count, 1),
            "employee_count": employee_count
        },
        "trends": {
            "vs_last_period": random.choice(["+12%", "-5%", "+3%", "-8%"]),
            "projected_next": random.choice(["Stable", "Slight increase expected", "Seasonal decrease expected"])
        }
    }

def tool_get_overtime_metrics(start, end, departments):
    """Generate realistic overtime metrics."""
    depts = (departments if departments and departments != ["all"] else ["HR", "Engineering", "Sales", "Marketing", "Finance", "Operations"])
    
    dept_hours = {}
    for d in depts:
        # Engineering and Operations typically have more OT
        base = {"Engineering": 400, "Operations": 350, "Sales": 200, "HR": 100, "Marketing": 150, "Finance": 180}.get(d.title(), 200)
        dept_hours[d.title()] = int(base * random.uniform(0.7, 1.3))
    
    total_hours = sum(dept_hours.values())
    employee_count = len(depts) * 25
    hourly_rate = 45  # Average OT rate
    
    return {
        "range": "Selected period" if not start else f"{start:%b %d} – {end:%b %d, %Y}",
        "total_hours": total_hours,
        "avg_per_employee": round(total_hours / employee_count, 1),
        "by_department": dept_hours,
        "est_cost": int(total_hours * hourly_rate * 1.5),  # 1.5x for OT rate
        "peak_day": random.choice(["Friday", "Thursday", "Wednesday"]),
        "compliance": {
            "exceeding_limit": random.randint(3, 12),  # Employees exceeding OT limits
            "risk_level": random.choice(["Low", "Moderate", "Elevated"])
        }
    }

def tool_get_attendance_metrics(start, end, departments):
    """Generate realistic attendance metrics."""
    attendance_rate = round(random.uniform(92, 97), 1)
    employee_count = len(departments) * 25 if departments else 150
    
    return {
        "range": "Selected period" if not start else f"{start:%b %d} – {end:%b %d, %Y}",
        "attendance_rate": attendance_rate,
        "unplanned_absences": random.randint(100, 200),
        "perfect_attendance": random.randint(50, 80),
        "avg_sick_days": round(random.uniform(2.5, 4.5), 1),
        "patterns": {
            "monday_absence_rate": round(random.uniform(5, 8), 1),
            "friday_absence_rate": round(random.uniform(6, 9), 1),
            "mid_week_rate": round(random.uniform(3, 5), 1)
        },
        "employee_count": employee_count
    }

def tool_get_balance_metrics(start, end, departments):
    """Generate realistic balance metrics."""
    employee_count = len(departments) * 25 if departments else 150
    
    return {
        "range": "As of " + datetime.now().strftime("%B %d, %Y"),
        "low_balance_count": random.randint(15, 35),
        "total_accrual_vacation_days": employee_count * random.randint(15, 25),
        "total_accrual_sick_days": employee_count * random.randint(8, 12),
        "avg_vacation_balance": round(random.uniform(12, 22), 1),
        "avg_sick_balance": round(random.uniform(5, 10), 1),
        "expiring_soon": {
            "30_days": random.randint(5, 15),
            "60_days": random.randint(10, 25),
            "90_days": random.randint(15, 40)
        },
        "liability": {
            "total_days": employee_count * 18,
            "estimated_value": f"${employee_count * 18 * 280:,}"  # Assuming $280/day
        }
    }

TOOLS = {
    "get_leave_metrics": tool_get_leave_metrics,
    "get_overtime_metrics": tool_get_overtime_metrics,
    "get_attendance_metrics": tool_get_attendance_metrics,
    "get_balance_metrics": tool_get_balance_metrics
}

def build_plan(intent: Intent) -> ExecutionPlan:
    """Build execution plan from Intent object."""
    rt = intent.report_type
    
    # Parse time range if available
    start, end = None, None
    if intent.time_range:
        try:
            start = datetime.fromisoformat(intent.time_range[0])
            end = datetime.fromisoformat(intent.time_range[1])
        except:
            pass

    departments = intent.departments or []
    steps: List[PlanStep] = []
    assumptions: List[str] = []

    if rt == "leave_analysis":
        steps.append(PlanStep("Leave Metrics", "get_leave_metrics", 
                            {"start": start, "end": end, "departments": departments}, "leave_metrics"))
        assumptions.append("Company holidays and weekends excluded from leave calculations.")
        assumptions.append("Pending leave requests not included in totals.")
    elif rt == "overtime_analysis":
        steps.append(PlanStep("Overtime Metrics", "get_overtime_metrics", 
                            {"start": start, "end": end, "departments": departments}, "ot_metrics"))
        assumptions.append("Overtime rate calculated at 1.5x standard hourly rate.")
        assumptions.append("Exempt employees excluded from overtime calculations.")
    elif rt == "attendance_analysis":
        steps.append(PlanStep("Attendance Metrics", "get_attendance_metrics", 
                            {"start": start, "end": end, "departments": departments}, "att_metrics"))
        assumptions.append("Weekends and holidays excluded from attendance calculations.")
        assumptions.append("Remote work days counted as present.")
    elif rt == "balance_report":
        steps.append(PlanStep("Balance Metrics", "get_balance_metrics", 
                            {"start": start, "end": end, "departments": departments}, "bal_metrics"))
        assumptions.append("Balances include approved but not yet taken leave.")
        assumptions.append("Accrual rates based on current policy as of report date.")
    else:
        # Default to comprehensive report
        steps.extend([
            PlanStep("Leave Metrics", "get_leave_metrics", 
                    {"start": start, "end": end, "departments": departments}, "leave_metrics"),
            PlanStep("Attendance Metrics", "get_attendance_metrics", 
                    {"start": start, "end": end, "departments": departments}, "att_metrics")
        ])
        assumptions.append("Comprehensive overview requested; showing multiple metric categories.")

    return ExecutionPlan(phase=Phase.PLAN, steps=steps, assumptions=assumptions)

# ============================
# Public service entry points
# ============================
async def analyze_report(payload: ReportAnalysisRequest, request: Request):
    """Analyze user query using LLM to extract intent."""
    rid = request.headers.get("X-Request-ID") or uuid.uuid4().hex
    q = (payload.query or "").strip()
    language = payload.language or "en-US"
    
    logger.info(f"rid={rid} reports.analyze start query='{payload.query}' lang={language}")

    # Use LLM client to analyze intent
    intent = llm_client.analyze_intent(q, preferred_lang=language)
    
    # Build execution plan from intent
    plan = build_plan(intent)
    
    # Convert Intent to dict for API response
    analysis = {
        "report_type": intent.report_type,
        "time_period": intent.time_period,
        "departments": intent.departments,
        "metrics": intent.metrics,
        "confidence": intent.confidence,
        "language": language
    }
    
    if intent.time_range:
        analysis["time_range"] = {
            "start": intent.time_range[0],
            "end": intent.time_range[1]
        }
    
    analysis["proposed_plan"] = {
        "phase": plan.phase.value,
        "steps": [{"name": s.name, "fn": s.fn, "args": {k: (v.isoformat() if isinstance(v, datetime) else v) 
                  for k, v in s.args.items()}} for s in plan.steps],
        "assumptions": plan.assumptions
    }

    result = {
        "needs_clarification": intent.needs_clarification,
        "analysis": analysis
    }
    
    if intent.needs_clarification:
        result["clarification_questions"] = intent.clarification_questions

    logger.info(f"rid={rid} reports.analyze complete type={analysis['report_type']} "
                f"confidence={intent.confidence:.2f} clarification={intent.needs_clarification}")
    
    return result

async def generate_report(payload: ReportGenerationRequest, request: Request):
    """Generate report with AI-powered narrative."""
    rid = request.headers.get("X-Request-ID") or uuid.uuid4().hex
    report_id = uuid.uuid4().hex
    language = payload.language or "en-US"
    
    logger.info(f"rid={rid} reports.generate start report_id={report_id} lang={language}")

    analysis = payload.analysis or {}
    
    # Reconstruct Intent from analysis
    intent = Intent(
        report_type=analysis.get("report_type", "unknown"),
        time_period=analysis.get("time_period", "monthly"),
        time_range=(analysis.get("time_range", {}).get("start"), 
                   analysis.get("time_range", {}).get("end")) if analysis.get("time_range") else None,
        departments=analysis.get("departments", []),
        metrics=analysis.get("metrics", []),
        confidence=analysis.get("confidence", 0.75)
    )

    # Generate title based on intent
    type_titles = {
        "leave_analysis": f"{intent.time_period.title()} Leave Analysis Report",
        "overtime_analysis": f"{intent.time_period.title()} Overtime Summary Report",
        "attendance_analysis": f"{intent.time_period.title()} Attendance Analysis Report",
        "balance_report": "Employee Balance Summary Report",
        "unknown": "Custom HR Analytics Report"
    }
    title = type_titles.get(intent.report_type, "HR Report")

    # Execute plan to gather data
    plan = build_plan(intent)
    data_bucket: Dict[str, Any] = {}
    
    for step in plan.steps:
        fn = TOOLS[step.fn]
        out = fn(**step.args)
        key = step.output_key or step.fn
        data_bucket[key] = out

    # Generate AI narrative
    narrative = llm_client.build_narrative(
        query=payload.query,
        intent=intent,
        data_bucket=data_bucket,
        title=title,
        target_language=language
    )

    try:
        # Generate the enhanced DOCX report
        doc_path = await _generate_enhanced_docx_report(
            report_id=report_id,
            title=narrative.title,
            query=payload.query,
            intent=intent,
            plan=plan,
            data_bucket=data_bucket,
            narrative=narrative,
            language=language
        )

        # Store session
        report_sessions[report_id] = {
            "title": narrative.title,
            "file_path": doc_path,
            "created_at": datetime.now(),
            "query": payload.query,
            "analysis": analysis,
            "language": language,
            "plan": {
                "phase": plan.phase.value,
                "steps": [{"name": s.name, "fn": s.fn} for s in plan.steps],
                "assumptions": plan.assumptions
            }
        }

        logger.info(f"rid={rid} reports.generate complete report_id={report_id} title='{narrative.title}'")
        
        return {
            "report_id": report_id,
            "title": narrative.title,
            "download_url": f"/api/reports/download/{report_id}",
            "created_at": datetime.now().isoformat(),
            "language": language
        }
        
    except Exception as e:
        logger.exception(f"rid={rid} reports.generate error: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

def download_report_response(report_id: str, request: Request):
    """Download generated report."""
    rid = request.headers.get("X-Request-ID") or uuid.uuid4().hex
    logger.info(f"rid={rid} reports.download start report_id={report_id}")

    if report_id not in report_sessions:
        logger.warning(f"rid={rid} reports.download report not found report_id={report_id}")
        raise HTTPException(status_code=404, detail="Report not found")

    session = report_sessions[report_id]
    file_path = session["file_path"]

    if not os.path.exists(file_path):
        logger.error(f"rid={rid} reports.download file not found path={file_path}")
        raise HTTPException(status_code=404, detail="Report file not found")

    filename = f"{session['title'].replace(' ', '_')}.docx"
    logger.info(f"rid={rid} reports.download serving file={filename}")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ===========================
# Enhanced DOCX builder
# ===========================
async def _generate_enhanced_docx_report(
    report_id: str,
    title: str,
    query: str,
    intent: Intent,
    plan: ExecutionPlan,
    data_bucket: Dict[str, Any],
    narrative: Narrative,
    language: str
) -> str:
    """Generate a professional, AI-enhanced DOCX report."""
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"report_{report_id}.docx")
    doc = Document()
    
    # Configure document styles
    _setup_document_styles(doc)

    # Title Page
    _add_title_page(doc, narrative.title, report_id)

    # Executive Summary
    _add_executive_summary(doc, narrative)

    # Methodology Section
    _add_methodology_section(doc, narrative, plan)

    # Key Insights
    _add_key_insights_section(doc, narrative)

    # Data Analysis Sections
    rt = intent.report_type
    if rt == "leave_analysis":
        _add_leave_analysis_section(doc, data_bucket.get("leave_metrics", {}))
    elif rt == "overtime_analysis":
        _add_overtime_analysis_section(doc, data_bucket.get("ot_metrics", {}))
    elif rt == "attendance_analysis":
        _add_attendance_analysis_section(doc, data_bucket.get("att_metrics", {}))
    elif rt == "balance_report":
        _add_balance_report_section(doc, data_bucket.get("bal_metrics", {}))
    else:
        # Add multiple sections for comprehensive report
        if "leave_metrics" in data_bucket:
            _add_leave_analysis_section(doc, data_bucket["leave_metrics"])
        if "att_metrics" in data_bucket:
            _add_attendance_analysis_section(doc, data_bucket["att_metrics"])

    # Risks & Mitigation
    _add_risks_section(doc, narrative)

    # Recommendations
    _add_recommendations_section(doc, narrative)

    # Appendix
    _add_appendix(doc, narrative, intent, query)

    doc.save(file_path)
    return file_path

def _setup_document_styles(doc: Document): #type: ignore
    """Configure professional document styles."""
    styles = doc.styles
    
    # Configure heading styles
    for i in range(1, 4):
        heading_style = styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if i == 1 else RGBColor(0x2E, 0x74, 0xB5)
        heading_style.font.size = Pt([24, 18, 14][i-1])
        heading_style.font.bold = True

    # Configure normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)

def _add_title_page(doc: Document, title: str, report_id: str): #type: ignore
    """Add professional title page."""
    # Add company logo placeholder
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(f"Report ID: {report_id}\n").bold = False
    meta_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n")
    meta_para.add_run("HR Analytics Department")
    
    # Page break
    doc.add_page_break()

def _add_executive_summary(doc: Document, narrative: Narrative): #type: ignore
    """Add executive summary section."""
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(narrative.executive_summary)
    doc.add_paragraph()

def _add_methodology_section(doc: Document, narrative: Narrative, plan: ExecutionPlan): #type: ignore
    """Add methodology section."""
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(narrative.methodology)
    
    if plan.assumptions:
        doc.add_heading('Key Assumptions', level=2)
        for assumption in plan.assumptions:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(assumption)
    
    doc.add_paragraph()

def _add_key_insights_section(doc: Document, narrative: Narrative): #type: ignore
    """Add key insights section with AI-generated insights."""
    doc.add_heading('Key Insights', level=1)
    
    for i, insight in enumerate(narrative.key_insights, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(insight)
    
    doc.add_paragraph()

def _add_leave_analysis_section(doc: Document, data: Dict[str, Any]): #type: ignore
    """Add enhanced leave analysis section."""
    doc.add_heading('Leave Analysis Details', level=1)
    doc.add_paragraph(f"Reporting Period: {data.get('range', 'Selected period')}")
    
    # Summary metrics table
    doc.add_heading('Summary Metrics', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    
    totals = data.get("totals", {})
    metrics = [
        ("Total Leave Days", str(totals.get("leave_days", "—"))),
        ("Average per Employee", f"{totals.get('avg_days_per_employee', '—')} days"),
        ("Total Employees", str(totals.get("employee_count", "—"))),
        ("Trend vs Last Period", data.get("trends", {}).get("vs_last_period", "—")),
    ]
    
    for metric, value in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
    
    doc.add_paragraph()
    
    # Leave by type
    doc.add_heading('Leave Distribution by Type', level=2)
    _add_dict_table(doc, "Leave Type", "Days", data.get("by_type", {}))
    
    doc.add_paragraph()
    
    # Department breakdown
    doc.add_heading('Department Analysis', level=2)
    _add_dict_table(doc, "Department", "Leave Days", data.get("by_department", {}))

def _add_overtime_analysis_section(doc: Document, data: Dict[str, Any]): #type: ignore
    """Add enhanced overtime analysis section."""
    doc.add_heading('Overtime Analysis Details', level=1)
    doc.add_paragraph(f"Reporting Period: {data.get('range', 'Selected period')}")
    
    # Key metrics
    doc.add_heading('Overtime Metrics', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    
    metrics = [
        ("Total Overtime Hours", f"{data.get('total_hours', '—'):,}"),
        ("Average per Employee", f"{data.get('avg_per_employee', '—')} hours"),
        ("Estimated Cost", f"${data.get('est_cost', 0):,}"),
        ("Peak Overtime Day", data.get("peak_day", "—")),
        ("Compliance Risk Level", data.get("compliance", {}).get("risk_level", "—")),
        ("Employees Exceeding Limits", str(data.get("compliance", {}).get("exceeding_limit", "—"))),
    ]
    
    for metric, value in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
    
    doc.add_paragraph()
    
    # Department breakdown
    doc.add_heading('Overtime by Department', level=2)
    _add_dict_table(doc, "Department", "Hours", data.get("by_department", {}))

def _add_attendance_analysis_section(doc: Document, data: Dict[str, Any]): #type: ignore
    """Add enhanced attendance analysis section."""
    doc.add_heading('Attendance Analysis Details', level=1)
    doc.add_paragraph(f"Reporting Period: {data.get('range', 'Selected period')}")
    
    # Summary metrics
    doc.add_heading('Attendance Metrics', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    
    metrics = [
        ("Overall Attendance Rate", f"{data.get('attendance_rate', '—')}%"),
        ("Unplanned Absences", str(data.get("unplanned_absences", "—"))),
        ("Perfect Attendance Count", str(data.get("perfect_attendance", "—"))),
        ("Average Sick Days", f"{data.get('avg_sick_days', '—')} days"),
        ("Total Employees", str(data.get("employee_count", "—"))),
    ]
    
    for metric, value in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
    
    doc.add_paragraph()
    
    # Patterns
    patterns = data.get("patterns", {})
    if patterns:
        doc.add_heading('Absence Patterns', level=2)
        doc.add_paragraph(f"Monday Absence Rate: {patterns.get('monday_absence_rate', '—')}%")
        doc.add_paragraph(f"Friday Absence Rate: {patterns.get('friday_absence_rate', '—')}%")
        doc.add_paragraph(f"Mid-Week Rate: {patterns.get('mid_week_rate', '—')}%")

def _add_balance_report_section(doc: Document, data: Dict[str, Any]): #type: ignore
    """Add enhanced balance report section."""
    doc.add_heading('Leave Balance Analysis', level=1)
    doc.add_paragraph(data.get('range', 'As of today'))
    
    # Balance metrics
    doc.add_heading('Balance Summary', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    
    metrics = [
        ("Low Balance Employees (<5 days)", str(data.get("low_balance_count", "—"))),
        ("Average Vacation Balance", f"{data.get('avg_vacation_balance', '—')} days"),
        ("Average Sick Leave Balance", f"{data.get('avg_sick_balance', '—')} days"),
        ("Total Vacation Accrual", f"{data.get('total_accrual_vacation_days', '—'):,} days"),
        ("Total Sick Leave Accrual", f"{data.get('total_accrual_sick_days', '—'):,} days"),
    ]
    
    for metric, value in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
    
    doc.add_paragraph()
    
    # Expiration warnings
    expiring = data.get("expiring_soon", {})
    if expiring:
        doc.add_heading('Expiration Alerts', level=2)
        doc.add_paragraph(f"Expiring in 30 days: {expiring.get('30_days', '—')} employees")
        doc.add_paragraph(f"Expiring in 60 days: {expiring.get('60_days', '—')} employees")
        doc.add_paragraph(f"Expiring in 90 days: {expiring.get('90_days', '—')} employees")
    
    # Liability
    liability = data.get("liability", {})
    if liability:
        doc.add_paragraph()
        doc.add_heading('Financial Liability', level=2)
        doc.add_paragraph(f"Total Accrued Days: {liability.get('total_days', '—'):,}")
        doc.add_paragraph(f"Estimated Value: {liability.get('estimated_value', '—')}")
 
def _add_risks_section(doc: Document, narrative: Narrative): #type: ignore
    """Add risks and mitigation section."""
    if narrative.risks:
        doc.add_heading('Risk Assessment', level=1)
        
        for i, risk in enumerate(narrative.risks, 1):
            p = doc.add_paragraph()
            p.add_run(f"Risk {i}: ").bold = True
            p.add_run(risk)
        
        doc.add_paragraph()

def _add_recommendations_section(doc: Document, narrative: Narrative): #type: ignore
    """Add recommendations section."""
    if narrative.recommendations:
        doc.add_heading('Recommendations', level=1)
        
        for i, rec in enumerate(narrative.recommendations, 1):
            p = doc.add_paragraph()
            # Use check mark for visual appeal
            p.add_run("✓ ").font.color.rgb = RGBColor(0x00, 0x70, 0x00)
            p.add_run(f"{rec}")
        
        doc.add_paragraph()

def _add_appendix(doc: Document, narrative: Narrative, intent: Intent, query: str): #type: ignore
    """Add appendix with technical details."""
    doc.add_page_break()
    doc.add_heading('Appendix', level=1)
    
    # Original Query
    doc.add_heading('Original Request', level=2)
    doc.add_paragraph(f'"{query}"')
    doc.add_paragraph()
    
    # Technical Details
    doc.add_heading('Technical Details', level=2)
    doc.add_paragraph(f"Report Type: {intent.report_type.replace('_', ' ').title()}")
    doc.add_paragraph(f"Time Period: {intent.time_period.title()}")
    doc.add_paragraph(f"Departments: {', '.join(intent.departments) if intent.departments else 'All'}")
    doc.add_paragraph(f"Confidence Score: {intent.confidence:.0%}")
    doc.add_paragraph()
    
    # Notes
    if narrative.appendix_notes:
        doc.add_heading('Additional Notes', level=2)
        for note in narrative.appendix_notes:
            doc.add_paragraph(f"• {note}", style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Generated by HCM AI Analytics Platform").italic = True
    footer.add_run(f"\n{datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

def _add_dict_table(doc: Document, key_header: str, value_header: str, data: Dict[str, Any]): #type: ignore
    """Helper to add a formatted table from dictionary data."""
    if not data:
        doc.add_paragraph("No data available for this period.")
        return
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'
    
    hdr = table.rows[0].cells
    hdr[0].text = key_header
    hdr[1].text = value_header
    
    # Sort by value for better readability
    sorted_items = sorted(data.items(), key=lambda x: x[1] if isinstance(x[1], (int, float)) else 0, reverse=True)
    
    for key, value in sorted_items:
        row = table.add_row().cells
        row[0].text = str(key)
        # Format numbers with commas
        if isinstance(value, (int, float)):
            row[1].text = f"{value:,}" if isinstance(value, int) else f"{value:,.1f}"
        else:
            row[1].text = str(value)